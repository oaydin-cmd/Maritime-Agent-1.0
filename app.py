import os
import shutil
import io
import datetime
import sqlite3
import streamlit as st

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

# HuggingFaceEmbeddings için esnek import bloğu (ImportError önleyici)
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.documents import Document

# Belge Yükleyiciler, OCR ve Word Oluşturma Araçları
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, UnstructuredExcelLoader
from pdf2image import convert_from_path
import pytesseract
import docx

os.environ["CURL_CA_BUNDLE"] = ""

# ---------------------------------------------------------
# 1. SAYFA YAPILANDIRMASI
# ---------------------------------------------------------
st.set_page_config(page_title="Denizcilik SMS & Asistan", page_icon="⚓", layout="wide")
st.title("⚓ Maritime Agent")

DOCS_DIR = "docs"
INDEX_DIR = "faiss_index"
DB_PATH = "chat_history.db"

# ---------------------------------------------------------
# WORD DOKÜMANI OLUŞTURMA YARDIMCI FONKSİYONU
# ---------------------------------------------------------
def create_docx_bytes(content_text, title="Denizcilik SMS Raporu"):
    doc = docx.Document()
    doc.add_heading(title, 0)
    
    # Metni satır satır ekle
    for line in content_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", "").strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", "").strip(), level=3)
        else:
            doc.add_paragraph(line)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ---------------------------------------------------------
# 2. SQLITE VERİTABANI
# ---------------------------------------------------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            role TEXT,
            content TEXT,
            docx_blob BLOB,
            file_name TEXT
        )
    ''')
    conn.commit()
    conn.close()

def save_message_to_db(role, content, docx_bytes=None, file_name=None):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    binary_data = sqlite3.Binary(docx_bytes) if isinstance(docx_bytes, bytes) else None
    cursor.execute(
        "INSERT INTO messages (timestamp, role, content, docx_blob, file_name) VALUES (?, ?, ?, ?, ?)",
        (now, role, content, binary_data, file_name)
    )
    conn.commit()
    conn.close()

def load_messages_from_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT id, timestamp, role, content, docx_blob, file_name FROM messages ORDER BY id ASC")
    rows = cursor.fetchall()
    conn.close()
    
    messages = []
    for row in rows:
        msg = {
            "id": row[0],
            "timestamp": row[1],
            "role": row[2],
            "content": row[3]
        }
        if row[4] and isinstance(row[4], (bytes, bytearray)) and len(row[4]) > 0:
            msg["docx_bytes"] = bytes(row[4])
            msg["file_name"] = row[5] or "SMS_Raporu.docx"
        messages.append(msg)
    return messages

def clear_db_history():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM messages")
    conn.commit()
    conn.close()

init_db()

# ---------------------------------------------------------
# 3. YALIN SOL MENÜ
# ---------------------------------------------------------
st.sidebar.header("⚙️ Sistem Ayarları")

api_key = st.secrets.get("OPENROUTER_API_KEY") or st.secrets.get("DEEPSEEK_API_KEY") or os.getenv("OPENROUTER_API_KEY") or os.getenv("DEEPSEEK_API_KEY")

if api_key:
    st.sidebar.success("🔑 API Key Tanımlı")
else:
    st.sidebar.error("❌ API Key Bulunamadı! Lütfen `.streamlit/secrets.toml` dosyanızı kontrol edin.")

st.sidebar.markdown("---")
st.sidebar.header("🧠 Hafıza Yönetimi")

if st.sidebar.button("🗑️ Tüm Sohbet Hafızasını Sıfırla"):
    clear_db_history()
    st.session_state.messages = []
    st.success("Hafıza temizlendi.")
    st.rerun()

# ---------------------------------------------------------
# 4. EMBEDDINGS VE DOKÜMAN İNDEKSLEME
# ---------------------------------------------------------
@st.cache_resource
def get_embeddings():
    return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

embeddings = get_embeddings()

def load_all_documents(folder_path):
    all_documents = []
    if not os.path.exists(folder_path):
        return all_documents, f"'{folder_path}' klasörü bulunamadı."

    for file in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file)
        ext = os.path.splitext(file)[1].lower()
        try:
            if ext == ".pdf":
                loader = PyPDFLoader(file_path)
                docs = loader.load()
                for d in docs:
                    d.metadata["source_file"] = file
                    d.metadata["page_label"] = f"Sayfa {d.metadata.get('page', 0) + 1}"
                if not docs or sum(len(d.page_content.strip()) for d in docs) < 50:
                    images = convert_from_path(file_path)
                    ocr_docs = []
                    for i, image in enumerate(images):
                        text = pytesseract.image_to_string(image, lang="tur+eng")
                        if text.strip():
                            ocr_docs.append(Document(page_content=text, metadata={"source_file": file, "page_label": f"Sayfa {i+1} (OCR)"}))
                    docs = ocr_docs
                all_documents.extend(docs)

            elif ext == ".docx":
                try:
                    loader = Docx2txtLoader(file_path)
                    docx_docs = loader.load()
                    for d in docx_docs:
                        d.metadata["source_file"] = file
                        d.metadata["page_label"] = "Word Dokümanı"
                    all_documents.extend(docx_docs)
                except Exception:
                    doc = docx.Document(file_path)
                    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    if full_text.strip():
                        all_documents.append(Document(page_content=full_text, metadata={"source_file": file, "page_label": "Word Dokümanı"}))

            elif ext == ".xlsx":
                loader = UnstructuredExcelLoader(file_path, mode="single")
                excel_docs = loader.load()
                for d in excel_docs:
                    d.metadata["source_file"] = file
                    d.metadata["page_label"] = "Excel Sayfası"
                all_documents.extend(excel_docs)
        except Exception as e:
            st.error(f"❌ '{file}' işlenirken hata: {str(e)}")

    return all_documents, None

@st.cache_resource
def load_or_create_vectorstore():
    if os.path.exists(INDEX_DIR):
        return FAISS.load_local(INDEX_DIR, embeddings, allow_dangerous_deserialization=True)
    documents, error = load_all_documents(DOCS_DIR)
    if error or not documents:
        return None
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    splits = text_splitter.split_documents(documents)
    vectorstore = FAISS.from_documents(splits, embeddings)
    vectorstore.save_local(INDEX_DIR)
    return vectorstore

vectorstore = load_or_create_vectorstore()

if st.sidebar.button("🔄 Doküman İndeksini Yenile"):
    if os.path.exists(INDEX_DIR):
        shutil.rmtree(INDEX_DIR)
    st.cache_resource.clear()
    st.rerun()

# ---------------------------------------------------------
# 5. AKILLI API VE DINAMIK PROMPT YAPISI
# ---------------------------------------------------------
retriever = None

if api_key and vectorstore:
    retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

    if api_key.startswith("sk-or-"):
        api_base = "https://openrouter.ai/api/v1"
        target_models = [
            "google/gemini-2.0-flash-001",
            "meta-llama/llama-3.3-70b-instruct",
            "deepseek/deepseek-chat"
        ]
    else:
        api_base = "https://api.deepseek.com"
        target_models = ["deepseek-chat", "deepseek-reasoner"]

    current_time_str = datetime.datetime.now().strftime("%d.%m.%Y, %A")

    system_prompt = (
        f"Bugünün tarihi ve günü: {current_time_str}.\n"
        "Sen akıllı, zeki ve doğal yanıt veren bir asistansın. Şirket dokümanları, formlar, raporlar ve denizcilik/SMS konularında uzmansın.\n\n"
        "KURALLAR:\n"
        "1. Kullanıcı ne derse ona doğrudan, mantıklı ve insan gibi yanıt ver.\n"
        "2. Tarih, gün veya saat sorulduğunda sana verilen güncel tarih bilgisini kullan.\n"
        "3. Form, liste veya rapor istendiğinde düzenli, başlıklandırılmış ve net bir format sun.\n"
        "4. 'REFERANS DOKÜMAN' alanında bilgi varsa gelen metindeki verileri analiz edip detaylıca yanıtla.\n\n"
        "REFERANS DOKÜMAN:\n{context}"
    )

    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{question}"),
    ])

# ---------------------------------------------------------
# 6. SOHBET ARAYÜZÜ VE İŞLEME
# ---------------------------------------------------------
if "messages" not in st.session_state:
    st.session_state.messages = load_messages_from_db()

for idx, message in enumerate(st.session_state.messages):
    with st.chat_message(message["role"]):
        st.markdown(message["content"])
        
        # Word İndirme Butonu Gösterimi
        if message.get("docx_bytes"):
            st.download_button(
                label="📥 Word Formatında İndir (.docx)",
                data=message["docx_bytes"],
                file_name=message.get("file_name", "SMS_Raporu.docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_{idx}"
            )

if user_input := st.chat_input("Mesajınızı yazın..."):
    if not api_key:
        st.error("⚠️ `.streamlit/secrets.toml` içinde geçerli bir API Key bulunamadı.")
    else:
        st.session_state.messages.append({"role": "user", "content": user_input})
        save_message_to_db("user", user_input)
        
        with st.chat_message("user"):
            st.markdown(user_input)

        with st.chat_message("assistant"):
            with st.spinner("Yazıyor..."):
                keywords = [
                    "sms", "prosedür", "denetim", "sire", "solas", "marpol", "kural", "form", 
                    "checklist", "tanker", "gemi", "güverte", "makine", "isps", "ism", "rapor", 
                    "manta", "pdf", "docx", "doküman", "belge", "dosya", "incele", "maddeler",
                    "hazırla", "oluştur", "özet", "analiz", "listele", "raporla"
                ]
                needs_rag = any(kw in user_input.lower() for kw in keywords)
                
                relevant_docs = []
                if vectorstore and needs_rag:
                    relevant_docs = retriever.invoke(user_input)
                
                context_text = "\n\n".join([d.page_content for d in relevant_docs]) if relevant_docs else "Yok"

                history_objs = []
                raw_history = load_messages_from_db()[:-1]
                for m in raw_history[-8:]:
                    if m["role"] == "user":
                        history_objs.append(HumanMessage(content=m["content"]))
                    else:
                        history_objs.append(AIMessage(content=m["content"]))

                formatted_prompt = prompt.format_messages(
                    context=context_text,
                    chat_history=history_objs,
                    question=user_input
                )
                
                llm_response = None
                last_error = ""

                for model_name in target_models:
                    try:
                        llm = ChatOpenAI(
                            model=model_name,
                            openai_api_key=api_key,
                            openai_api_base=api_base,
                            temperature=0.4,
                            timeout=15,
                            max_retries=0,
                            default_headers={"HTTP-Referer": "http://localhost:8501", "X-Title": "Expert Assistant"}
                        )
                        llm_response = llm.invoke(formatted_prompt)
                        break
                    except Exception as err:
                        last_error = str(err)
                        continue

                if llm_response:
                    response_text = llm_response.content if hasattr(llm_response, 'content') else str(llm_response)

                    sources = []
                    seen = set()
                    for doc in relevant_docs:
                        file_name = doc.metadata.get("source_file", doc.metadata.get("source", "Doküman"))
                        page_info = doc.metadata.get("page_label", "")
                        source_str = f"📄 {file_name} ({page_info})"
                        if source_str not in seen:
                            seen.add(source_str)
                            sources.append(source_str)

                    final_response = response_text
                    if sources:
                        final_response += "\n\n---\n**İncelediğim Şirket Dokümanları:**\n" + "\n".join([f"- {src}" for src in sources])

                    st.markdown(final_response)

                    # Form/Rapor taleplerinde veya genel yanıtlarda otomatik Word Dosyası (.docx) Oluşturma
                    form_keywords = ["form", "rapor", "hazırla", "oluştur", "docx", "word", "maddeler", "checklist", "incele"]
                    docx_bytes = None
                    file_name = None

                    if any(kw in user_input.lower() for kw in form_keywords):
                        file_name = f"SMS_Rapor_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                        docx_bytes = create_docx_bytes(final_response, title="SMS & Denizcilik Asistan Raporu")
                        
                        st.download_button(
                            label="📥 Word Formatında İndir (.docx)",
                            data=docx_bytes,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_new_{len(st.session_state.messages)}"
                        )

                    save_message_to_db("assistant", final_response, docx_bytes=docx_bytes, file_name=file_name)
                    st.session_state.messages.append({
                        "role": "assistant", 
                        "content": final_response,
                        "docx_bytes": docx_bytes,
                        "file_name": file_name
                    })
                else:
                    st.error(f"❌ API Hatası: {last_error if last_error else 'Servis yanıt vermedi.'}\n\nLütfen secrets dosyasındaki API key'inizi kontrol edin.")
